# -*- coding: utf-8 -*-
"""
Step9 自包含 HTML 报告生成。

把各 step 产出的图表（base64 内嵌，单文件可分享）+ 关键指标表 + 配置摘要，
渲染成一个不依赖任何外部资源的 report.html。可选导出 docx。
"""
from __future__ import annotations

import base64
import csv
import html
import json
from pathlib import Path

from .utils import get_logger, dump_json

log = get_logger()

# 数据文件 -> 中文标题（如存在则渲染为表格）
_TABLE_FILES = [
    ("phase_stats.csv", "各传播阶段规模"),
    ("phase_polarity.csv", "阶段情感极性占比(%)"),
    ("stance_compare.csv", "立场对比"),
    ("stance_dist.csv", "立场总体分布"),
    ("topic_summary.csv", "主题汇总"),
    ("topic_level_crosstab.csv", "主题 × 层级交叉"),
    ("network_metrics.json", "网络整体指标"),
    ("subnetwork_metrics.csv", "十大视频子网络"),
    ("word_cooccur_centrality.csv", "核心短语共现中心性"),
    ("cluster_explore.json", "无监督聚类探索"),
]


def _img_b64(path: str) -> str | None:
    p = Path(path)
    if not p.exists():
        return None
    return base64.b64encode(p.read_bytes()).decode("ascii")


def _csv_table(path: Path, max_rows=200) -> str:
    rows = []
    with path.open(encoding="utf-8", errors="ignore", newline="") as fh:
        r = csv.reader(fh)
        for i, row in enumerate(r):
            if i > max_rows:
                break
            rows.append(row)
    if not rows:
        return ""
    head, body = rows[0], rows[1:]
    th = "".join(f"<th>{html.escape(str(c))}</th>" for c in head)
    trs = []
    for row in body:
        tds = "".join(f"<td>{html.escape(str(c))}</td>" for c in row)
        trs.append(f"<tr>{tds}</tr>")
    return (f"<table class='data'><thead><tr>{th}</tr></thead>"
            f"<tbody>{''.join(trs)}</tbody></table>")


def _json_table(path: Path) -> str:
    obj = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(obj, dict):
        return f"<pre>{html.escape(json.dumps(obj, ensure_ascii=False, indent=1))}</pre>"
    rows = []
    for k, v in obj.items():
        if isinstance(v, (dict, list)):
            v = html.escape(json.dumps(v, ensure_ascii=False, indent=1))
        else:
            v = html.escape(str(v))
        rows.append(f"<tr><th>{html.escape(str(k))}</th><td>{v}</td></tr>")
    return f"<table class='data'>{''.join(rows)}</table>"


def _stats_block(stats: dict) -> str:
    if not stats:
        return ""
    rows = []
    for k, v in stats.items():
        if isinstance(v, dict):
            inner = "<br>".join(f"{html.escape(str(kk))}: {html.escape(str(vv))}"
                                for kk, vv in v.items())
            rows.append(f"<tr><th>{html.escape(str(k))}</th><td>{inner}</td></tr>")
        else:
            rows.append(f"<tr><th>{html.escape(str(k))}</th><td>{html.escape(str(v))}</td></tr>")
    return f"<table class='stats'>{''.join(rows)}</table>"


_CSS = """
* { box-sizing: border-box; }
body { font-family: -apple-system, "Segoe UI", "Microsoft YaHei", sans-serif;
       margin: 0; color: #222; background: #f6f7f9; }
header { background: #1f2d3d; color: #fff; padding: 28px 40px; }
header h1 { margin: 0 0 6px; font-size: 26px; }
header .meta { opacity: .8; font-size: 13px; }
main { max-width: 1080px; margin: 0 auto; padding: 24px 20px 60px; }
section { background: #fff; border-radius: 10px; padding: 20px 24px; margin: 18px 0;
          box-shadow: 0 1px 3px rgba(0,0,0,.08); }
section h2 { border-left: 4px solid #2471A3; padding-left: 10px; margin-top: 4px; }
figure { text-align: center; margin: 16px 0; }
figure img { max-width: 100%; border: 1px solid #eee; border-radius: 6px; }
figcaption { color: #666; font-size: 13px; margin-top: 6px; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13px; }
table.data th, table.data td { border: 1px solid #e2e6ea; padding: 5px 8px; text-align: left; }
table.data thead { background: #eef3f8; }
table.stats { width: auto; }
table.stats th { text-align: right; background: #f3f6fa; padding: 4px 10px; }
table.stats td { padding: 4px 10px; }
.toc a { color: #2471A3; text-decoration: none; margin-right: 14px; font-size: 14px; }
footer { text-align: center; color: #999; font-size: 12px; padding: 20px; }
"""


def build_report(cfg, results: list[dict], extra_tables: list[str] | None = None) -> str:
    out = cfg.out_root
    topic = cfg.topic
    figs = [(p, t, c) for r in results for (p, t, c) in r.get("figures", [])]

    # TOC
    toc = "".join(f"<a href='#sec{i}'>图{i+1} {html.escape(t)}</a>"
                  for i, (p, t, c) in enumerate(figs)) or "<span>本报告无图表</span>"

    # body: each figure as a figure block
    body = []
    for i, (p, t, c) in enumerate(figs, 1):
        b64 = _img_b64(p)
        if not b64:
            continue
        body.append(
            f"<section id='sec{i}'><h2>图{i} {html.escape(t)}</h2>"
            f"<figure><img src='data:image/png;base64,{b64}' alt='{html.escape(t)}'>"
            f"<figcaption>{html.escape(c or '')}</figcaption></figure></section>"
        )

    # 数据表
    tables_html = []
    for fname, title in _TABLE_FILES:
        fp = out / fname
        if not fp.exists():
            continue
        if fname.endswith(".json"):
            tbl = _json_table(fp)
        else:
            tbl = _csv_table(fp)
        if tbl:
            tables_html.append(f"<section><h2>{html.escape(title)}</h2>{tbl}</section>")
    if extra_tables:
        for fp in extra_tables:
            p = Path(fp)
            if p.exists() and p.suffix in (".csv", ".json"):
                tbl = _csv_table(p) if p.suffix == ".csv" else _json_table(p)
                if tbl:
                    tables_html.append(f"<section><h2>{html.escape(p.stem)}</h2>{tbl}</section>")

    # 配置摘要
    try:
        eff = (out / "effective_config.yaml").read_text(encoding="utf-8")
    except Exception:
        eff = ""
    config_html = (f"<section><h2>运行配置（effective_config.yaml）</h2>"
                   f"<pre class='cfg'>{html.escape(eff)}</pre></section>") if eff else ""

    # 关键指标汇总
    stats_all = []
    for r in results:
        if r.get("stats"):
            stats_all.append(_stats_block(r["stats"]))
    stats_html = "".join(stats_all) if stats_all else ""

    html_doc = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(topic)} · B 站舆情分析报告</title>
<style>{_CSS}</style></head>
<body>
<header><h1>{html.escape(topic)} · B 站舆情演化分析报告</h1>
<div class="meta">由 BiliOpinion 生成 · 输出目录：{html.escape(str(out))}</div></header>
<main>
<section><h2>目录</h2><div class="toc">{toc}</div></section>
{('<section><h2>关键指标汇总</h2>' + stats_html + '</section>') if stats_html else ''}
{''.join(body)}
{''.join(tables_html)}
{config_html}
</main>
<footer>BiliOpinion · 计算传播学 B 站舆情分析工具 · 自包含报告</footer>
</body></html>"""

    report_path = out / "report.html"
    report_path.write_text(html_doc, encoding="utf-8")
    log.info("HTML 报告已生成: %s（%d 张图，%.0f KB）",
             report_path, len(figs), len(html_doc) / 1024)
    return str(report_path)


def build_docx(cfg, results: list[dict]) -> str | None:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as e:  # noqa: BLE001
        log.warning("未安装 python-docx，跳过 docx 报告：%s", e)
        return None
    out = cfg.out_root
    doc = Document()
    doc.add_heading(f"{cfg.topic} · B 站舆情演化分析报告", level=0)
    for r in results:
        for (p, t, c) in r.get("figures", []):
            fp = Path(p)
            if not fp.exists():
                continue
            doc.add_heading(t, level=1)
            try:
                doc.add_picture(str(fp), width=Inches(6.0))
            except Exception:
                pass
            if c:
                doc.add_paragraph(c)
    path = out / "report.docx"
    doc.save(str(path))
    log.info("docx 报告已生成: %s", path)
    return str(path)
